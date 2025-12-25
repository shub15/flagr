import io
import re
import fitz  # PyMuPDF
import spacy
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load NLP model efficiently (do this once at startup)
# Ensure you run: python -m spacy download en_core_web_sm
nlp = spacy.load("en_core_web_sm")

def get_sensitive_terms(text: str) -> list[str]:
    """
    Combines NLP (Name/Org) and Regex (Email/Phone) to find redaction terms.
    """
    redaction_set = set()

    # 1. Regex for Patterns (Email, Phone, generic ID patterns)
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    redaction_set.update(re.findall(email_pattern, text))
    redaction_set.update(re.findall(phone_pattern, text))

    # 2. NLP for Entities (Companies, Persons)
    # Limit text length for NLP to avoid memory spikes on huge docs
    doc = nlp(text[:100000]) 
    for ent in doc.ents:
        if ent.label_ in ["ORG", "PERSON", "GPE"]: # GPE = Geopolitical Entity (Countries, Cities)
            redaction_set.add(ent.text)

    # Filter out short/common words to avoid over-redaction (e.g., "The", "it")
    final_terms = [term for term in redaction_set if len(term) > 2]
    logger.info(f"Identified sensitive terms: {final_terms}")
    return final_terms

def process_redaction_pdf(file_bytes: bytes) -> io.BytesIO:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    
    # 1. Extract text to find terms
    full_text = "".join([page.get_text() for page in doc])
    logger.info(f"PDF Extraction Preview (First 500 chars): {full_text[:500]}")
    
    terms = get_sensitive_terms(full_text)

    # 2. Redact
    total_redactions = 0
    for page in doc:
        for term in terms:
            text_instances = page.search_for(term)
            for inst in text_instances:
                page.add_redact_annot(inst, fill=(0, 0, 0)) # Black box
                total_redactions += 1
        page.apply_redactions()
    
    logger.info(f"Total redactions applied: {total_redactions}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def process_redaction_docx(file_bytes: bytes) -> io.BytesIO:
    doc = Document(io.BytesIO(file_bytes))
    
    # 1. Extract text
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)
    
    combined_text = "\n".join(full_text)
    logger.info(f"DOCX Extraction Preview (First 500 chars): {combined_text[:500]}")
    terms = get_sensitive_terms(combined_text)

    def replace_text(text, terms):
        for term in terms:
            if term in text:
                text = text.replace(term, "█" * len(term)) # Visual redaction
        return text

    # 2. Redact Paragraphs
    for para in doc.paragraphs:
        para.text = replace_text(para.text, terms)

    # 3. Redact Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = replace_text(para.text, terms)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

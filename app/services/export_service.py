"""
Export service for generating Word redlines and PDF reports from contract reviews.
"""

import logging
from typing import List
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from app.models.schemas import ContractReviewResult, ReviewCategory

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting contract reviews to various formats."""
    
    def generate_word_redline(
        self,
        review: ContractReviewResult,
        contract_text: str,
        output_path: str
    ) -> str:
        """
        Generate a Word document with redline markups.
        
        Args:
            review: Contract review result
            contract_text: Original contract text
            output_path: Path to save the DOCX file
        
        Returns:
            Path to generated file
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Contract Review - Redline Markup', 0)
            title.alignment = 1  # Center
            
            # Summary
            doc.add_heading('Review Summary', level=1)
            summary_table = doc.add_table(rows=3, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            
            cells = summary_table.rows[0].cells
            cells[0].text = 'Review ID'
            cells[1].text = review.review_id
            
            cells = summary_table.rows[1].cells
            cells[0].text = 'Safety Score'
            cells[1].text = f"{review.safety_score}/100"
            
            cells = summary_table.rows[2].cells
            cells[0].text = 'Total Findings'
            cells[1].text = str(review.total_findings)
            
            doc.add_paragraph()
            
            # Findings sections
            if review.critical_points:
                doc.add_heading('CRITICAL Issues', level=1)
                for i, point in enumerate(review.critical_points, 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(f"[CRITICAL] ").bold = True
                    p.add_run(point.advice)
                    if point.quote:
                        quote_p = doc.add_paragraph(style='Quote')
                        quote_p.add_run(f'Quote: "{point.quote}"')
                        # Highlight quote in red
                        quote_run = quote_p.runs[0]
                        quote_run.font.color.rgb = RGBColor(255, 0, 0)
                    if point.legal_reference:
                        ref_p = doc.add_paragraph(style='Intense Quote')
                        ref_p.add_run(f'Legal Reference: {point.legal_reference}')
                doc.add_paragraph()
            
            if review.missing_points:
                doc.add_heading('MISSING Clauses', level=1)
                for i, point in enumerate(review.missing_points, 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(f"[MISSING] ").bold = True
                    p.add_run(point.advice)
                    if point.legal_reference:
                        ref_p = doc.add_paragraph(style='Intense Quote')
                        ref_p.add_run(f'Legal Reference: {point.legal_reference}')
                doc.add_paragraph()
            
            if review.negotiable_points:
                doc.add_heading('NEGOTIABLE Terms', level=1)
                for i, point in enumerate(review.negotiable_points, 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(f"[NEGOTIABLE] ").bold = True
                    p.add_run(point.advice)
                    if point.quote:
                        quote_p = doc.add_paragraph(style='Quote')
                        quote_p.add_run(f'Quote: "{point.quote}"')
                        # Highlight in orange
                        quote_run = quote_p.runs[0]
                        quote_run.font.color.rgb = RGBColor(255, 165, 0)
                doc.add_paragraph()
            
            if review.good_points:
                doc.add_heading('GOOD Points (Compliant)', level=1)
                for i, point in enumerate(review.good_points, 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(f"[GOOD] ").bold = True
                    p.add_run(point.advice)
                    if point.quote:
                        quote_p = doc.add_paragraph(style='Quote')
                        quote_p.add_run(f'Quote: "{point.quote}"')
                        # Highlight in green
                        quote_run = quote_p.runs[0]
                        quote_run.font.color.rgb = RGBColor(0, 128, 0)
                doc.add_paragraph()
            
            # Original contract with annotations
            doc.add_page_break()
            doc.add_heading('Original Contract with Annotations', level=1)
            
            # Add contract text with simple paragraph breaks
            paragraphs = contract_text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Save
            doc.save(output_path)
            logger.info(f"Generated Word redline: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate Word redline: {e}")
            raise
    
    def generate_pdf_report(
        self,
        review: ContractReviewResult,
        output_path: str
    ) -> str:
        """
        Generate a PDF summary report.
        
        Args:
            review: Contract review result
            output_path: Path to save the PDF file
        
        Returns:
            Path to generated file
        """
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=12,
                spaceBefore=12
            )
            
            # Title
            elements.append(Paragraph("Contract Review Report", title_style))
            elements.append(Spacer(1, 12))
            
            # Summary Table
            summary_data = [
                ['Review ID', review.review_id],
                ['Date', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')],
                ['Safety Score', f'{review.safety_score}/100'],
                ['Total Findings', str(review.total_findings)],
                ['Critical Issues', str(len(review.critical_points))],
                ['Missing Clauses', str(len(review.missing_points))],
                ['Negotiable Terms', str(len(review.negotiable_points))],
                ['Good Points', str(len(review.good_points))]
            ]
            
            summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8f0f7')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey)
            ]))
            
            elements.append(summary_table)
            elements.append(Spacer(1, 20))
            
            # Add findings sections
            def add_findings_section(title: str, findings: List, color: str):
                if not findings:
                    return
                
                elements.append(Paragraph(title, heading_style))
                elements.append(Spacer(1, 6))
                
                for i, point in enumerate(findings, 1):
                    # Finding text
                    text = f"<b>{i}.</b> {point.advice}"
                    elements.append(Paragraph(text, styles['Normal']))
                    
                    # Quote if available
                    if point.quote:
                        quote_text = f'<i>Contract Quote: "{point.quote[:200]}{"..." if len(point.quote) > 200 else ""}"</i>'
                        elements.append(Paragraph(quote_text, styles['Normal']))
                    
                    # Legal reference if available
                    if point.legal_reference:
                        ref_text = f'<font color="{color}"><b>Legal Reference:</b> {point.legal_reference[:200]}</font>'
                        elements.append(Paragraph(ref_text, styles['Normal']))
                    
                    # Confidence
                    conf_text = f'<font size=8>Agent: {point.agent_source} | Confidence: {point.confidence:.0%}</font>'
                    elements.append(Paragraph(conf_text, styles['Normal']))
                    elements.append(Spacer(1, 10))
                
                elements.append(Spacer(1, 12))
            
            # Add each section
            add_findings_section('CRITICAL Issues ⚠️', review.critical_points, '#ff0000')
            add_findings_section('MISSING Clauses 📋', review.missing_points, '#ff8800')
            add_findings_section('NEGOTIABLE Terms 💼', review.negotiable_points, '#0088ff')
            add_findings_section('GOOD Points ✅', review.good_points, '#00aa00')
            
            # Footer note
            elements.append(Spacer(1, 30))
            disclaimer = Paragraph(
                '<i>Note: This is an automated review and does not constitute legal advice. '
                'Please consult with a qualified attorney for legal guidance.</i>',
                styles['Normal']
            )
            elements.append(disclaimer)
            
            # Build PDF
            doc.build(elements)
            logger.info(f"Generated PDF report: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate PDF report: {e}")
            raise


# Global instance
export_service = ExportService()
